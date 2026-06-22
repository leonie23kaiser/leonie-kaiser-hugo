import re, sys, pathlib
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEAL=RGBColor(0x08,0x65,0x84); TEAL6=RGBColor(0x5F,0xA2,0xA0)
VIOLET=RGBColor(0x6B,0x2C,0x8C); INK=RGBColor(0x1d,0x22,0x28); MUTE=RGBColor(0x5b,0x65,0x72)

def add_runs(p, text):
    # inline bold (**), italic (*), code (`)
    parts = re.split(r'(\*\*.+?\*\*|\*.+?\*|`.+?`)', text)
    for part in parts:
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            r=p.add_run(part[2:-2]); r.bold=True
        elif part.startswith('*') and part.endswith('*'):
            r=p.add_run(part[1:-1]); r.italic=True
        elif part.startswith('`') and part.endswith('`'):
            r=p.add_run(part[1:-1]); r.font.name='Consolas'
        else:
            p.add_run(part)

def md2docx(src, dst):
    lines = pathlib.Path(src).read_text('utf-8').splitlines()
    doc = Document()
    n=doc.styles['Normal']; n.font.name='Calibri'; n.font.size=Pt(11); n.font.color.rgb=INK
    i=0
    while i < len(lines):
        line = lines[i].rstrip()
        # table block
        if '|' in line and i+1 < len(lines) and re.match(r'^\s*\|?[\s:|-]+\|?\s*$', lines[i+1]):
            header=[c.strip() for c in line.strip().strip('|').split('|')]
            i+=2; rows=[]
            while i < len(lines) and '|' in lines[i]:
                rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')]); i+=1
            t=doc.add_table(rows=1, cols=len(header)); t.style='Light Grid Accent 1'
            for j,h in enumerate(header):
                cell=t.rows[0].cells[j]; cell.paragraphs[0].clear()
                add_runs(cell.paragraphs[0], h)
                for rr in cell.paragraphs[0].runs: rr.bold=True
            for row in rows:
                cells=t.add_row().cells
                for j,val in enumerate(row):
                    if j<len(cells):
                        cells[j].paragraphs[0].clear(); add_runs(cells[j].paragraphs[0], val)
            doc.add_paragraph(); continue
        if line.startswith('#### '):
            p=doc.add_paragraph(); r=p.add_run(line[5:]); r.bold=True; r.font.color.rgb=VIOLET; r.font.size=Pt(11)
        elif line.startswith('### '):
            p=doc.add_paragraph(); r=p.add_run(line[4:]); r.bold=True; r.font.color.rgb=TEAL6; r.font.size=Pt(13)
        elif line.startswith('## '):
            p=doc.add_paragraph(); r=p.add_run(line[3:]); r.bold=True; r.font.color.rgb=TEAL; r.font.size=Pt(16)
        elif line.startswith('# '):
            p=doc.add_paragraph(); r=p.add_run(line[2:]); r.bold=True; r.font.color.rgb=TEAL; r.font.size=Pt(22)
        elif line.startswith('> '):
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.3)
            add_runs(p, line[2:]); 
            for r in p.runs: r.italic=True; r.font.color.rgb=MUTE
        elif re.match(r'^\s*[-*] ', line):
            indent = (len(line)-len(line.lstrip()))//2
            p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.left_indent=Inches(0.25+0.25*indent)
            add_runs(p, re.sub(r'^\s*[-*] ','',line))
        elif re.match(r'^\s*\d+\. ', line):
            p=doc.add_paragraph(style='List Number'); add_runs(p, re.sub(r'^\s*\d+\. ','',line))
        elif line.strip()=='---':
            pass
        elif line.strip()=='':
            doc.add_paragraph()
        else:
            p=doc.add_paragraph(); add_runs(p, line)
        i+=1
    doc.save(dst); print("saved", dst)

md2docx(sys.argv[1], sys.argv[2])
