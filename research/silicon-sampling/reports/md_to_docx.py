#!/usr/bin/env python3
"""Konvertiert reports/REPORT.md -> Word (.docx) mit Brand-Styling.
Behandelt: H1/H2/H3, Blockquote, Bullets, Tabellen, O-Töne mit <sub>-Quellen,
Inline **fett**/*kursiv*/`code`, Trennlinien."""
import re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEAL = RGBColor(0x08,0x65,0x84); VIOLET = RGBColor(0x6B,0x2C,0x8C)
INK = RGBColor(0x1D,0x22,0x28); MUTE = RGBColor(0x5B,0x65,0x72); GOLD=RGBColor(0xCF,0x98,0x2B)

SRC, OUT = sys.argv[1], sys.argv[2]
lines = open(SRC, encoding="utf-8").read().split("\n")

doc = Document()
# Basis-Stil
st = doc.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(10.5); st.font.color.rgb = INK
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Cm(2.0); sec.left_margin = sec.right_margin = Cm(2.0)
for hname, sz, col in [("Heading 1",16,TEAL),("Heading 2",13,TEAL),("Heading 3",11.5,VIOLET)]:
    h = doc.styles[hname]; h.font.name="Arial"; h.font.size=Pt(sz); h.font.bold=True; h.font.color.rgb=col

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")

def shade(cell, hexc):
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),hexc)
    cell._tc.get_or_add_tcPr().append(sh)

def add_runs(p, text, base_color=INK, base_size=None, base_italic=False):
    text = text.replace("<sub>","").replace("</sub>","")
    for part in INLINE.split(text):
        if not part: continue
        r = p.add_run(); r.font.color.rgb = base_color; r.italic = base_italic
        if base_size: r.font.size = Pt(base_size)
        if part.startswith("**") and part.endswith("**"):
            r.text = part[2:-2]; r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r.text = part[1:-1]; r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r.text = part[1:-1]; r.font.name="Consolas"; r.font.color.rgb=MUTE
        else:
            r.text = part

def is_table(i):
    return lines[i].startswith("|") and i+1 < len(lines) and re.match(r"\|[\s:|-]+\|", lines[i+1])

i = 0
while i < len(lines):
    ln = lines[i]
    s = ln.strip()
    if not s:
        i += 1; continue
    # Tabellen
    if ln.startswith("|") and is_table(i):
        header = [c.strip() for c in ln.strip("|").split("|")]
        rows = []
        j = i + 2
        while j < len(lines) and lines[j].startswith("|"):
            rows.append([c.strip() for c in lines[j].strip("|").split("|")]); j += 1
        tbl = doc.add_table(rows=1, cols=len(header)); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.style = "Table Grid"
        for k,c in enumerate(header):
            cell = tbl.rows[0].cells[k]; cell.paragraphs[0].clear(); shade(cell, "086584")
            add_runs(cell.paragraphs[0], c);
            for rr in cell.paragraphs[0].runs: rr.bold=True; rr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        for ri,row in enumerate(rows):
            cells = tbl.add_row().cells
            for k,c in enumerate(row):
                if k < len(cells):
                    cells[k].paragraphs[0].clear(); add_runs(cells[k].paragraphs[0], c)
                    if ri % 2 == 1: shade(cells[k], "EDF7F2")
        doc.add_paragraph()
        i = j; continue
    # Überschriften
    if s.startswith("# "):
        doc.add_heading(s[2:], level=1); i+=1; continue
    if s.startswith("## "):
        doc.add_heading(s[3:], level=2); i+=1; continue
    if s.startswith("### "):
        doc.add_heading(s[4:], level=3); i+=1; continue
    # Trennlinie
    if s == "---":
        p = doc.add_paragraph(); pr = p._p.get_or_add_pPr(); pb = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        for a,v in [("w:val","single"),("w:sz","6"),("w:space","1"),("w:color","ECE3D8")]: bot.set(qn(a),v)
        pb.append(bot); pr.append(pb); i+=1; continue
    # Blockquote (ESOMAR-Hinweis etc.)
    if s.startswith(">"):
        block = []
        while i < len(lines) and lines[i].strip().startswith(">"):
            block.append(lines[i].strip()[1:].strip()); i += 1
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
        pr = p._p.get_or_add_pPr(); pb = OxmlElement("w:pBdr"); lft = OxmlElement("w:left")
        for a,v in [("w:val","single"),("w:sz","18"),("w:space","8"),("w:color","CF982B")]: lft.set(qn(a),v)
        pb.append(lft); pr.append(pb)
        add_runs(p, " ".join(block))
        continue
    # O-Ton-Quelle (eingerückte <sub>-Zeile) -> an vorigen Absatz hängen
    if ln.startswith("  ") and "<sub>" in ln:
        p = doc.paragraphs[-1]; p.add_run("\n")
        add_runs(p, s, base_color=MUTE, base_size=8.5, base_italic=True); i+=1; continue
    # Bullets
    if s.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        # O-Ton-Bullet „...": leichte Hervorhebung
        add_runs(p, s[2:]); i+=1; continue
    # normaler Absatz
    p = doc.add_paragraph(); add_runs(p, s); i+=1

doc.save(OUT)
print("OK ->", OUT)
